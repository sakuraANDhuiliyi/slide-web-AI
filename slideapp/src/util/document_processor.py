import os
import io
import docx
import PyPDF2
# import textract  # 移除textract依赖
from django.conf import settings
import google.generativeai as genai
import tempfile

def extract_text_from_docx(file_path):
    """
    从Word文档中提取文本内容
    
    Args:
        file_path: 文档路径
        
    Returns:
        str: 提取的文本内容
    """
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # 提取标题
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                full_text.append(f"{'#' * int(para.style.name[-1])} {para.text}")
            else:
                full_text.append(para.text)
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text for cell in row.cells])
                full_text.append(f"| {row_text} |")
        
        return '\n'.join(full_text)
    except Exception as e:
        print(f"提取Word文档内容时出错: {str(e)}")
        # 不使用textract作为备选方法
        try:
            # 简单读取文件内容
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as inner_e:
            print(f"使用备选方法提取内容时出错: {str(inner_e)}")
            return ""

def extract_text_from_pdf(file_path):
    """
    从PDF文档中提取文本内容
    
    Args:
        file_path: PDF文档路径
        
    Returns:
        str: 提取的文本内容
    """
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = []
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text.append(page.extract_text())
                
            return '\n'.join(text)
    except Exception as e:
        print(f"提取PDF文档内容时出错: {str(e)}")
        # 不使用textract作为备选方法
        try:
            # 尝试二进制读取文件内容后转换为文本
            with open(file_path, 'rb') as f:
                content = f.read()
                try:
                    return content.decode('utf-8', errors='ignore')
                except:
                    return str(content)
        except Exception as inner_e:
            print(f"使用备选方法提取PDF内容时出错: {str(inner_e)}")
            return ""

def extract_text_from_file(file_path, file_type=None):
    """
    根据文件类型提取内容
    
    Args:
        file_path: 文件路径
        file_type: 文件类型（可选，如果不提供则根据扩展名判断）
        
    Returns:
        str: 提取的文本内容
    """
    if not file_type:
        _, file_extension = os.path.splitext(file_path)
        file_extension = file_extension.lower()
        
        if file_extension == '.docx' or file_extension == '.doc':
            file_type = 'word'
        elif file_extension == '.pdf':
            file_type = 'pdf'
        else:
            # 尝试使用通用方法
            file_type = 'other'
    
    if file_type == 'word':
        return extract_text_from_docx(file_path)
    elif file_type == 'pdf':
        return extract_text_from_pdf(file_path)
    else:
        # 尝试作为纯文本文件读取
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            try:
                # 如果纯文本读取失败，尝试二进制读取
                with open(file_path, 'rb') as f:
                    content = f.read()
                    return content.decode('utf-8', errors='ignore')
            except Exception as inner_e:
                print(f"使用通用方法提取内容时出错: {str(inner_e)}")
                return ""

def summarize_with_ai(text, title, slide_count=5, layout="horizontal_vertical"):
    """
    使用AI对文档内容进行总结，并生成幻灯片内容
    
    Args:
        text: 提取的文本内容
        title: 幻灯片标题
        slide_count: 期望生成的幻灯片数量
        layout: 幻灯片布局类型
    
    Returns:
        str: Markdown格式的幻灯片内容
    """
    try:
        print(f"[进度] 2/5 - 开始AI分析文档内容...")
        
        # 获取API密钥
        api_key = os.environ.get('GEMINI_API_KEY', '')
        if not api_key:
            api_key = getattr(settings, 'GEMINI_API_KEY', '')
        
        if not api_key:
            return None
        
        # 重置模型配置
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-2.0-flash')
        
        # 根据布局类型选择提示信息
        layout_instructions = ""
        if layout == "horizontal_vertical":
            layout_instructions = """
            4. 幻灯片布局规则：
               - 使用 "---" (三个减号的单独行) 来划分水平幻灯片（不同章节）
               - 使用 "----" (四个减号的单独行) 来划分垂直幻灯片（同一章节内的不同内容）
            5. 请严格遵循这种结构：
               
               # 第一章节
               内容
               
               ---
               
               # 第二章节
               
               ----
               
               ## 第二章节的第一个子页
               内容
               
               ----
               
               ## 第二章节的第二个子页
               内容
            """
        elif layout == "horizontal_only":
            layout_instructions = """
            4. 幻灯片布局规则：
               - 使用 "---" (三个减号的单独行) 来划分幻灯片
            5. 请严格遵循这种结构：
               
               # 第一个幻灯片
               内容
               
               ---
               
               # 第二个幻灯片
               内容
            """
        else:  # with_animation
            layout_instructions = """
            4. 幻灯片布局规则：
               - 使用 "---" (三个减号的单独行) 来划分水平幻灯片（不同章节）
               - 使用 "----" (四个减号的单独行) 来划分垂直幻灯片（同一章节内的不同内容）
               - 使用 "++++" (四个加号的单独行) 来创建渐变幻灯片（展示内容的变化过程）
            """
        
        # 构建提示信息
        prompt = f"""
        请将以下文档内容总结为一个Markdown格式的幻灯片，遵循以下要求：
        
        - 幻灯片标题: {title}
        - 幻灯片数量: 约{slide_count}张
        
        请使用以下Markdown语法格式生成幻灯片:
        
        1. 使用 # 后跟文字表示章节标题幻灯片
        2. 使用 ## 后跟文字表示普通幻灯片的标题
        3. 正文内容直接使用Markdown语法
        {layout_instructions}
        
        注意：
        - 每个分隔符必须独占一行，前后都有空行
        - 必须是纯Markdown格式，不要加额外的HTML或其他标记
        - 请基于文档原始内容进行总结，保留重要的知识点和结构
        - 提炼文档的主要观点、关键信息和核心内容
        
        文档内容:
        {text[:30000]}  # 限制输入长度
        """
        
        # 添加重试机制
        max_retries = 3
        retry_count = 0
        
        while retry_count < max_retries:
            try:
                # 生成内容
                print(f"[进度] 3/5 - 开始生成幻灯片结构...")
                response = model.generate_content(prompt)
                return response.text
            except Exception as e:
                retry_count += 1
                print(f"AI总结内容失败 (尝试 {retry_count}/{max_retries}): {str(e)}")
                if retry_count >= max_retries:
                    raise
                # 短暂延迟后重试
                import time
                time.sleep(1)
        
        return None
    except Exception as e:
        print(f"AI总结文档内容出错: {str(e)}")
        return None

def process_uploaded_document(uploaded_file, title, slide_count=5, layout="horizontal_vertical"):
    """
    处理上传的文档，提取内容并生成幻灯片
    
    Args:
        uploaded_file: 上传的文件对象
        title: 幻灯片标题
        slide_count: 期望生成的幻灯片数量
        layout: 幻灯片布局类型
        
    Returns:
        str: Markdown格式的幻灯片内容，失败则返回None
    """
    try:
        # 创建临时文件
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
            # 写入上传的文件内容
            for chunk in uploaded_file.chunks():
                temp_file.write(chunk)
            temp_file_path = temp_file.name
        
        try:
            # 提取文档内容
            text_content = extract_text_from_file(temp_file_path)
            
            if not text_content or len(text_content.strip()) < 50:
                return None, "无法从文档中提取内容或内容过少"
            
            # 使用AI总结内容生成幻灯片
            markdown_content = summarize_with_ai(text_content, title, slide_count, layout)
            
            return markdown_content, None
        finally:
            # 删除临时文件
            os.unlink(temp_file_path)
    
    except Exception as e:
        import traceback
        error_detail = ''.join(traceback.format_exception(type(e), e, e.__traceback__))
        print(f"处理上传文档出错: {error_detail}")
        return None, f"处理文档时出错: {str(e)}" 